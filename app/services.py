import os
import sys

import app
from app.constants import CODIGOS_DICTADO


class TranslationService:
    def traducir(self, texto, destino):
        from deep_translator import GoogleTranslator
        LIMITE = 4500
        parrafos = texto.split("\n")
        partes = []
        lote = ""
        for parrafo in parrafos:
            if len(lote) + len(parrafo) + 1 > LIMITE:
                if lote:
                    t = GoogleTranslator(source="auto", target=destino).translate(lote.strip())
                    partes.append(t)
                lote = parrafo + "\n"
            else:
                lote += parrafo + "\n"
        if lote.strip():
            t = GoogleTranslator(source="auto", target=destino).translate(lote.strip())
            partes.append(t)
        return "\n".join(partes)


class OcrService:
    def __init__(self):
        self._reader = None

    def extraer_texto(self, ruta_imagen, idioma_ui="Español (Argentina)"):
        if not app.OCR_DISPONIBLE:
            raise RuntimeError("easyocr no instalado")
        if self._reader is None:
            codigo = CODIGOS_DICTADO.get(idioma_ui, "es-AR")
            langs = [codigo.split("-")[0]]
            if langs[0] != "es":
                langs.append("es")
            import easyocr
            self._reader = easyocr.Reader(langs, gpu=False)
        resultado = self._reader.readtext(ruta_imagen, paragraph=True)
        return "\n".join(p[1] for p in resultado)


class FileService:
    def guardar(self, ruta, contenido):
        with open(ruta, "w", encoding="utf-8") as f:
            f.write(contenido)

    def abrir(self, ruta):
        with open(ruta, "r", encoding="utf-8") as f:
            return f.read()

    def exportar_docx(self, ruta, contenido, tamano_fuente=12):
        if not app.DOCX_DISPONIBLE:
            raise RuntimeError("python-docx no instalado")
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        for parrafo in contenido.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(parrafo if parrafo else " ")
            run.font.size = Pt(tamano_fuente)
        doc.save(ruta)

    def exportar_pdf(self, ruta, contenido, tamano_fuente=12):
        if not app.PDF_DISPONIBLE:
            raise RuntimeError("fpdf2 no instalado")
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        font_path = self._obtener_fuente_pdf()
        if font_path and os.path.isfile(font_path[0]):
            try:
                pdf.add_font("DejaVu", "", font_path[0])
                pdf.add_font("DejaVu", "B", font_path[1])
                pdf.set_font("DejaVu", "", tamano_fuente)
            except Exception as e:
                print(f"[PDF] Error al cargar DejaVu, usando Helvetica: {e}", file=sys.stderr)
                pdf.set_font("Helvetica", "", tamano_fuente)
        else:
            print("[PDF] Fuente DejaVuSans no encontrada. Usando Helvetica (sin acentos).", file=sys.stderr)
            pdf.set_font("Helvetica", "", tamano_fuente)
        for linea in contenido.split("\n"):
            if linea.strip():
                pdf.multi_cell(0, 6, linea)
            else:
                pdf.ln(4)
        pdf.output(ruta)

    def _obtener_fuente_pdf(self):
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        normal = os.path.join(base_dir, "DejaVuSans.ttf")
        bold = os.path.join(base_dir, "DejaVuSans-Bold.ttf")
        if os.path.isfile(normal) and os.path.isfile(bold):
            return (normal, bold)
        return (None, None)
